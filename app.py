from flask import Flask, render_template, request, redirect, url_for, send_file, abort, jsonify
from flask_sqlalchemy import SQLAlchemy
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime, timezone
import os
import re
import io
import zipfile
import tempfile
import uuid
from werkzeug.utils import secure_filename
import logging
from dateutil.parser import parse  # Requires: pip install python-dateutil
import json
from copy import deepcopy
import subprocess
import platform

# Initialize Flask app
app = Flask(__name__)

# Configuration for PythonAnywhere (replace 'mytypist' with your actual username)
BASE_DIR = os.path.abspath(os.path.dirname(__file__))

# BASE_DIR = '/home/mytypist/mytypist_app'
app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{os.path.join(BASE_DIR, "db", "db.sqlite")}'
app.config['UPLOAD_FOLDER'] = os.path.join(BASE_DIR, 'uploads')
app.config['GENERATED_FOLDER'] = os.path.join(BASE_DIR, 'generated')
app.config['ADMIN_KEY'] = os.environ.get('ADMIN_KEY', 'secretkey123')  # Set this in PythonAnywhere Web tab
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# Initialize database
db = SQLAlchemy(app)

# Ensure directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['GENERATED_FOLDER'], exist_ok=True)
os.makedirs(os.path.join(BASE_DIR, 'db'), exist_ok=True)

# **Database Models**
class Template(db.Model):
    __tablename__ = 'template'
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    type = db.Column(db.String(50), nullable=False)
    file_path = db.Column(db.String(200), nullable=False)
    font_family = db.Column(db.String(50), nullable=False)
    font_size = db.Column(db.Integer, nullable=False)
    is_active = db.Column(db.Boolean, default=True)
    placeholders = db.relationship('Placeholder', back_populates='template', cascade="all, delete-orphan")
    created_documents = db.relationship('CreatedDocument', back_populates='template', cascade="all, delete-orphan")

class Placeholder(db.Model):
    __tablename__ = 'placeholder'
    id = db.Column(db.Integer, primary_key=True)
    template_id = db.Column(db.Integer, db.ForeignKey('template.id'), nullable=False)
    name = db.Column(db.String(50), nullable=False)
    paragraph_index = db.Column(db.Integer, nullable=False)
    start_run_index = db.Column(db.Integer, nullable=False)
    end_run_index = db.Column(db.Integer, nullable=False)
    bold = db.Column(db.Boolean, default=False)
    italic = db.Column(db.Boolean, default=False)
    underline = db.Column(db.Boolean, default=False)
    casing = db.Column(db.String(20), default="none")
    template = db.relationship('Template', back_populates='placeholders')

class CreatedDocument(db.Model):
    __tablename__ = 'created_document'
    id = db.Column(db.Integer, primary_key=True)
    template_id = db.Column(db.Integer, db.ForeignKey('template.id'), nullable=False)
    user_name = db.Column(db.String(100), nullable=False)
    file_path = db.Column(db.String(200), nullable=False)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    batch_id = db.Column(db.String(50), nullable=True)  # For batch processing
    template = db.relationship('Template', back_populates='created_documents')

class BatchGeneration(db.Model):
    __tablename__ = 'batch_generation'
    id = db.Column(db.Integer, primary_key=True)
    batch_id = db.Column(db.String(50), unique=True, nullable=False)
    user_name = db.Column(db.String(100), nullable=False)
    template_ids = db.Column(db.Text, nullable=False)  # JSON list of template IDs
    user_inputs = db.Column(db.Text, nullable=False)  # JSON of user inputs
    zip_file_path = db.Column(db.String(200), nullable=True)
    status = db.Column(db.String(20), default='pending')  # pending, completed, failed
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    completed_at = db.Column(db.DateTime, nullable=True)

# **Helper Functions**
def ordinal(n):
    """Convert a number to its ordinal form (e.g., 1 -> 1st, 2 -> 2nd)."""
    if 11 <= (n % 100) <= 13:
        suffix = 'th'
    else:
        suffix = ['th', 'st', 'nd', 'rd', 'th'][min(n % 10, 4)]
    return str(n) + suffix

def format_date(date_string, template_type):
    """Format a date string based on the template type."""
    try:
        date_obj = parse(date_string)
        day = ordinal(date_obj.day)
        month = date_obj.strftime("%B")
        year = date_obj.year
        if template_type == "letter":
            return f"{day} {month}, {year}"
        elif template_type == "affidavit":
            return f"{day} of {month}, {year}"
        return f"{date_obj.day} {month} {year}"
    except ValueError:
        logger.warning(f"Invalid date format: {date_string}")
        return date_string

def extract_placeholders(doc):
    """Extract placeholders like ${name} from a Word document with enhanced robustness."""
    placeholders = []
    placeholder_pattern = re.compile(r'\$\{([^}]+)\}')
    
    # Process paragraphs
    for p_idx, paragraph in enumerate(doc.paragraphs):
        full_text = ''.join(run.text for run in paragraph.runs)
        matches = placeholder_pattern.finditer(full_text)
        for match in matches:
            placeholder_name = match.group(1).strip()
            start_pos = match.start()
            end_pos = match.end()
            current_pos = 0
            start_run_idx = end_run_idx = None
            bold = italic = underline = False
            font_name = None
            font_size = None
            
            for r_idx, run in enumerate(paragraph.runs):
                run_start = current_pos
                run_end = current_pos + len(run.text)
                if start_run_idx is None and run_start <= start_pos < run_end:
                    start_run_idx = r_idx
                    bold = run.font.bold or False
                    italic = run.font.italic or False
                    underline = run.font.underline or False
                    font_name = run.font.name
                    font_size = run.font.size.pt if run.font.size else None
                if run_start < end_pos <= run_end:
                    end_run_idx = r_idx
                    break
                current_pos = run_end
                
            if start_run_idx is not None and end_run_idx is not None:
                placeholders.append({
                    'paragraph_index': p_idx,
                    'start_run_index': start_run_idx,
                    'end_run_index': end_run_idx,
                    'name': placeholder_name,
                    'bold': bold,
                    'italic': italic,
                    'underline': underline,
                    'casing': 'none',
                    'font_name': font_name,
                    'font_size': font_size
                })
    
    # Process tables
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for p_idx, paragraph in enumerate(cell.paragraphs):
                    full_text = ''.join(run.text for run in paragraph.runs)
                    matches = placeholder_pattern.finditer(full_text)
                    for match in matches:
                        placeholder_name = match.group(1).strip()
                        placeholders.append({
                            'paragraph_index': -1,  # Special marker for table cells
                            'table_row': row_idx,
                            'table_cell': cell_idx,
                            'table_paragraph': p_idx,
                            'start_run_index': 0,
                            'end_run_index': 0,
                            'name': placeholder_name,
                            'bold': False,
                            'italic': False,
                            'underline': False,
                            'casing': 'none',
                            'font_name': None,
                            'font_size': None
                        })
    
    return placeholders

def detect_document_font(doc):
    """Detect the most common font and size in a document."""
    font_counts = {}
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name and run.font.size:
                key = (run.font.name, int(run.font.size.pt))
                font_counts[key] = font_counts.get(key, 0) + 1
    if font_counts:
        return max(font_counts.items(), key=lambda x: x[1])[0]
    return "Times New Roman", 12

def allowed_file(filename):
    """Check if a file has a .docx extension."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() == 'docx'

def set_default_font(doc, font_name, font_size):
    """Set the default font for a document."""
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)

def remove_empty_runs(doc):
    """Remove empty runs from a document to clean up formatting."""
    for para in doc.paragraphs:
        p = para._element
        runs = list(p.findall('.//w:r', namespaces=p.nsmap))
        for run in runs:
            t = run.find('.//w:t', namespaces=run.nsmap)
            if t is not None and t.text == '':
                p.remove(run)

def enhance_document_formatting(doc):
    """Apply enhanced formatting to improve document quality."""
    # Set up styles for better appearance
    styles = doc.styles
    
    # Enhance normal style
    try:
        normal_style = styles['Normal']
        normal_font = normal_style.font
        if not normal_font.name:
            normal_font.name = 'Times New Roman'
        if not normal_font.size:
            normal_font.size = Pt(12)
        
        # Set paragraph formatting
        normal_paragraph_format = normal_style.paragraph_format
        normal_paragraph_format.space_after = Pt(6)
        normal_paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    except KeyError:
        pass  # Style doesn't exist
    
    # Apply consistent formatting to all paragraphs
    for paragraph in doc.paragraphs:
        # Ensure consistent paragraph spacing
        if paragraph.paragraph_format.space_after is None:
            paragraph.paragraph_format.space_after = Pt(6)
        if paragraph.paragraph_format.line_spacing_rule is None:
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # Ensure all runs have consistent formatting
        for run in paragraph.runs:
            if run.font.name is None:
                run.font.name = 'Times New Roman'
            if run.font.size is None:
                run.font.size = Pt(12)

def create_enhanced_document(template_path, user_inputs, template):
    """Create a document with enhanced formatting and placeholder replacement."""
    # Load the template
    doc = Document(template_path)
    
    # Set default font
    set_default_font(doc, template.font_family, template.font_size)
    
    # Get all placeholders for this template
    placeholders = Placeholder.query.filter_by(template_id=template.id).order_by(
        Placeholder.paragraph_index, Placeholder.start_run_index).all()
    
    # Process placeholders
    for placeholder in placeholders:
        try:
            # Handle table placeholders
            if placeholder.paragraph_index == -1:
                # This is a table placeholder - handle separately
                continue
                
            if placeholder.paragraph_index >= len(doc.paragraphs):
                logger.warning(f"Invalid paragraph index {placeholder.paragraph_index} for placeholder {placeholder.name}")
                continue
                
            paragraph = doc.paragraphs[placeholder.paragraph_index]
            
            if (placeholder.start_run_index >= len(paragraph.runs) or 
                placeholder.end_run_index >= len(paragraph.runs)):
                logger.warning(f"Invalid run indices for placeholder {placeholder.name}")
                continue

            user_input = user_inputs.get(placeholder.name, "")
            formatted_text = process_placeholder_text(user_input, placeholder, template)
            
            # Apply the replacement
            apply_placeholder_replacement(paragraph, placeholder, formatted_text, template)
            
        except Exception as e:
            logger.error(f"Error processing placeholder {placeholder.name}: {str(e)}")
            continue
    
    # Clean up the document
    remove_empty_runs(doc)
    enhance_document_formatting(doc)
    
    return doc

def process_placeholder_text(user_input, placeholder, template):
    """Process placeholder text with enhanced formatting rules."""
    if not user_input:
        return ""
    
    formatted_text = user_input.strip()
    
    # Date formatting
    if "date" in placeholder.name.lower() or "date_ofbirth" in placeholder.name.lower():
        formatted_text = format_date(user_input, template.type)
    
    # Address formatting for letters
    elif "address" in placeholder.name.lower() and template.type == "letter":
        # Keep address formatting as is - handled separately
        pass
    
    # Apply casing
    elif placeholder.casing == "upper":
        formatted_text = formatted_text.upper()
    elif placeholder.casing == "lower":
        formatted_text = formatted_text.lower()
    elif placeholder.casing == "title":
        formatted_text = formatted_text.title()
    
    return formatted_text

def apply_placeholder_replacement(paragraph, placeholder, formatted_text, template):
    """Apply placeholder replacement with preserved formatting."""
    # Special handling for addresses in letters
    if ("address" in placeholder.name.lower() and template.type == "letter" and 
        "," in formatted_text):
        parts = [part.strip() for part in formatted_text.split(",")]
        if parts:
            # Clear the target runs
            if placeholder.start_run_index != placeholder.end_run_index:
                for r_idx in range(placeholder.start_run_index + 1, placeholder.end_run_index + 1):
                    if r_idx < len(paragraph.runs):
                        paragraph.runs[r_idx].text = ""
            
            run = paragraph.runs[placeholder.start_run_index]
            run.clear()
            
            # Add address parts with line breaks
            for i, part in enumerate(parts):
                run.add_text(part.strip())
                if i < len(parts) - 1:
                    run.add_text(",")
                    run.add_break()
                elif not part.endswith("."):
                    run.add_text(".")
            
            # Apply formatting
            run.font.name = template.font_family
            run.font.size = Pt(template.font_size)
            run.bold = placeholder.bold
            run.italic = placeholder.italic
            run.underline = placeholder.underline
    else:
        # Standard replacement
        run = paragraph.runs[placeholder.start_run_index]
        
        # Clear multiple runs if needed
        if placeholder.start_run_index != placeholder.end_run_index:
            for r_idx in range(placeholder.start_run_index + 1, placeholder.end_run_index + 1):
                if r_idx < len(paragraph.runs):
                    paragraph.runs[r_idx].text = ""
        
        # Set the text
        run.text = formatted_text
        
        # Apply formatting
        run.font.name = template.font_family
        run.font.size = Pt(template.font_size)
        run.bold = placeholder.bold
        run.italic = placeholder.italic
        run.underline = placeholder.underline
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Improve paragraph spacing and alignment
    for para in doc.paragraphs:
        # Skip empty paragraphs
        if not para.text.strip():
            continue
            
        # Improve spacing
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        
        # Set line spacing
        para.paragraph_format.line_spacing = 1.15
        
        # Align headings and specific content based on template type
        if template_type == "letter":
            # For letters, align date to right
            if any(word in para.text.lower() for word in ["date:", "dated:"]):
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Align signature blocks to right
            elif any(word in para.text.lower() for word in ["sincerely", "regards", "yours", "faithfully"]):
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        elif template_type == "affidavit":
            # Center title for affidavits
            if "affidavit" in para.text.lower() or "declaration" in para.text.lower():
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    
        # Apply general formatting improvements
        if len(para.text) < 50 and para.text.isupper():
            # Likely a heading
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(12)
            
    return doc

def add_page_numbers(doc):
    """Add page numbers to the document footer."""
    for i, section in enumerate(doc.sections):
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        # Add styling to page numbers
        run.font.size = Pt(9)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    return doc

def convert_docx_to_pdf(docx_path, pdf_path):
    """Convert DOCX file to PDF using LibreOffice or similar."""
    try:
        if platform.system() == "Windows":
            # Try using LibreOffice if available
            try:
                subprocess.run([
                    "soffice", "--headless", "--convert-to", "pdf", 
                    "--outdir", os.path.dirname(pdf_path), docx_path
                ], check=True, capture_output=True)
                return True
            except (subprocess.CalledProcessError, FileNotFoundError):
                # LibreOffice not available, try alternative method
                logger.warning("LibreOffice not found, PDF conversion not available")
                return False
        else:
            # Linux/Unix systems
            try:
                subprocess.run([
                    "libreoffice", "--headless", "--convert-to", "pdf",
                    "--outdir", os.path.dirname(pdf_path), docx_path
                ], check=True, capture_output=True)
                return True
            except (subprocess.CalledProcessError, FileNotFoundError):
                logger.warning("LibreOffice not found, PDF conversion not available")
                return False
    except Exception as e:
        logger.error(f"PDF conversion failed: {str(e)}")
        return False

# **Routes**
@app.route('/')
def index():
    """Display the homepage with template types and recent documents."""
    # Get all active templates and filter out ones with missing files
    all_templates = Template.query.filter_by(is_active=True).all()
    valid_types = set()
    
    for template in all_templates:
        template_file_path = os.path.join(app.config['UPLOAD_FOLDER'], template.file_path)
        if os.path.exists(template_file_path):
            valid_types.add(template.type)
    
    types = sorted(list(valid_types))
    page = int(request.args.get('page', 1))
    per_page = 10
    recent_docs = CreatedDocument.query.filter(CreatedDocument.template.has(is_active=True))\
        .order_by(CreatedDocument.created_at.desc())\
        .offset((page-1)*per_page).limit(per_page).all()
    total_docs = CreatedDocument.query.filter(CreatedDocument.template.has(is_active=True)).count()
    total_pages = (total_docs + per_page - 1) // per_page
    return render_template('index.html', types=types, recent_docs=recent_docs,
                         page=page, total_pages=total_pages, admin_key=app.config['ADMIN_KEY'])

@app.route('/templates')
def get_templates():
    """Return a JSON list of templates for a given type."""
    type_ = request.args.get('type')
    if not type_:
        return jsonify({'error': 'Type is required'}), 400
    
    # Get all active templates for this type and filter out ones with missing files
    all_templates = Template.query.filter_by(type=type_, is_active=True).all()
    valid_templates = []
    
    for template in all_templates:
        template_file_path = os.path.join(app.config['UPLOAD_FOLDER'], template.file_path)
        if os.path.exists(template_file_path):
            valid_templates.append({'id': template.id, 'name': template.name})
    
    return jsonify(valid_templates)

@app.route('/batch')
def batch_selection():
    """Display the batch document generation page."""
    # Get all active templates and filter out ones with missing files
    all_templates = Template.query.filter_by(is_active=True).all()
    valid_templates = []
    valid_types = set()
    
    for template in all_templates:
        template_file_path = os.path.join(app.config['UPLOAD_FOLDER'], template.file_path)
        if os.path.exists(template_file_path):
            valid_templates.append(template)
            valid_types.add(template.type)
    
    return render_template('batch.html', types=sorted(list(valid_types)), templates=valid_templates)

@app.route('/batch-placeholders')
def get_batch_placeholders():
    """Return combined placeholders for selected templates."""
    template_ids = request.args.getlist('template_ids[]')
    if not template_ids:
        return jsonify({'error': 'No templates selected'}), 400
    
    combined_placeholders = set()
    templates_info = []
    
    for template_id in template_ids:
        template = Template.query.filter_by(id=template_id, is_active=True).first()
        if template:
            # Check if template file exists
            template_file_path = os.path.join(app.config['UPLOAD_FOLDER'], template.file_path)
            if os.path.exists(template_file_path):
                placeholders = Placeholder.query.filter_by(template_id=template_id).all()
                template_placeholders = [ph.name for ph in placeholders]
                combined_placeholders.update(template_placeholders)
                templates_info.append({
                    'id': template.id,
                    'name': template.name,
                    'type': template.type,
                    'placeholders': template_placeholders
                })
    
    return jsonify({
        'placeholders': sorted(list(combined_placeholders)),
        'templates': templates_info
    })

@app.route('/create/<int:template_id>')
def create(template_id):
    """Render the document creation page for a specific template."""
    template = Template.query.filter_by(id=template_id, is_active=True).first_or_404()
    placeholders = Placeholder.query.filter_by(template_id=template_id)\
        .order_by(Placeholder.paragraph_index, Placeholder.start_run_index).all()
    unique_names = []
    seen = set()
    for ph in placeholders:
        if ph.name not in seen:
            unique_names.append(ph.name)
            seen.add(ph.name)
    return render_template('create.html', template=template, placeholder_names=unique_names)

@app.route('/generate', methods=['POST'])
def generate():
    """Generate a document from a template and user inputs."""
    template_id = request.form['template_id']
    template = Template.query.filter_by(id=template_id, is_active=True).first_or_404()
    user_inputs = {key: request.form[key] for key in request.form if key != 'template_id' and key != 'batch_mode'}

    # Check if template file exists
    template_file_path = os.path.join(app.config['UPLOAD_FOLDER'], template.file_path)
    if not os.path.exists(template_file_path):
        logger.error(f"Template file not found: {template_file_path}")
        return render_template('error.html', message=f"Template file not found: {template.name}"), 404

    try:
        doc = Document(template_file_path)
    except Exception as e:
        logger.error(f"Error loading template {template.name}: {str(e)}")
        return render_template('error.html', message="Failed to load template. Please contact administrator."), 500
    set_default_font(doc, template.font_family, template.font_size)
    placeholders = Placeholder.query.filter_by(template_id=template.id)\
        .order_by(Placeholder.paragraph_index, Placeholder.start_run_index).all()

    for placeholder in placeholders:
        paragraph = doc.paragraphs[placeholder.paragraph_index]
        if placeholder.start_run_index >= len(paragraph.runs) or placeholder.end_run_index >= len(paragraph.runs):
            logger.warning(f"Invalid run indices for placeholder {placeholder.name} in paragraph {placeholder.paragraph_index}")
            continue

        user_input = user_inputs.get(placeholder.name, "")
        formatted_text = user_input

        if "date" in placeholder.name.lower() or "date_ofbirth" in placeholder.name.lower():
            formatted_text = format_date(user_input, template.type)

        elif "address" in placeholder.name.lower() and template.type == "letter":
            parts = [part.strip() for part in user_input.split(",")]
            if parts:
                if placeholder.start_run_index != placeholder.end_run_index:
                    for r_idx in range(placeholder.start_run_index + 1, placeholder.end_run_index + 1):
                        paragraph.runs[r_idx].text = ""
                run = paragraph.runs[placeholder.start_run_index]
                run.clear()
                for i, part in enumerate(parts):
                    run.add_text(part)
                    if i == len(parts) - 1 or part.endswith("."):
                        if not part.endswith("."):
                            run.add_text(".")
                        break
                    else:
                        run.add_text(",")
                        run.add_break()
                run.font.name = template.font_family
                run.font.size = Pt(template.font_size)
                run.bold = placeholder.bold
                run.italic = placeholder.italic
                run.underline = placeholder.underline
                continue

        else:
            if placeholder.casing == "upper":
                formatted_text = formatted_text.upper()
            elif placeholder.casing == "lower":
                formatted_text = formatted_text.lower()
            elif placeholder.casing == "title":
                formatted_text = formatted_text.title()

        run = paragraph.runs[placeholder.start_run_index]
        if placeholder.start_run_index == placeholder.end_run_index:
            run.text = formatted_text
        else:
            logger.debug(f"Placeholder {placeholder.name} spans multiple runs ({placeholder.start_run_index} to {placeholder.end_run_index})")
            for r_idx in range(placeholder.start_run_index + 1, placeholder.end_run_index + 1):
                paragraph.runs[r_idx].text = ""
            run.text = formatted_text
        run.font.name = template.font_family
        run.font.size = Pt(template.font_size)
        run.bold = placeholder.bold
        run.italic = placeholder.italic
        run.underline = placeholder.underline

    # Apply enhanced formatting
    remove_empty_runs(doc)
    enhance_document_formatting(doc, template.type)
    add_page_numbers(doc)

    user_name = user_inputs.get("name", "Unknown").strip()
    user_name = re.sub(r'\s+', '_', user_name)
    template_name = template.name.strip()
    template_name = re.sub(r'\s+', '_', template_name)
    current_date = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    file_name = f"{user_name}_{template_name}_{current_date}.docx"
    file_path = os.path.join(app.config['GENERATED_FOLDER'], file_name)
    doc.save(file_path)

    created_doc = CreatedDocument(template_id=template.id, user_name=user_name, file_path=file_name)
    db.session.add(created_doc)
    db.session.commit()
    
    # Redirect to results page to show download options
    return redirect(url_for('show_results', doc_id=created_doc.id))

@app.route('/batch-generate', methods=['POST'])
def batch_generate():
    """Generate multiple documents from selected templates."""
    # Get template IDs from form data
    template_ids_str = request.form.get('template_ids', '')
    template_ids = [tid.strip() for tid in template_ids_str.split(',') if tid.strip()]
    
    if not template_ids:
        template_ids = request.form.getlist('template_ids[]')
    
    user_inputs = {key: request.form[key] for key in request.form if key not in ['template_ids'] and not key.startswith('template_ids')}
    
    if not template_ids:
        return render_template('error.html', message='No templates selected'), 400
    
    # Generate individual documents
    generated_docs = []
    successful_generations = 0
    batch_id = str(uuid.uuid4())
    
    for template_id in template_ids:
        template = Template.query.filter_by(id=template_id, is_active=True).first()
        if not template:
            continue
        
        # Check if template file exists
        template_file_path = os.path.join(app.config['UPLOAD_FOLDER'], template.file_path)
        if not os.path.exists(template_file_path):
            logger.warning(f"Template file not found: {template_file_path}")
            continue
            
        try:
            doc = Document(template_file_path)
        except Exception as e:
            logger.error(f"Error loading template {template.name}: {str(e)}")
            continue
        
        set_default_font(doc, template.font_family, template.font_size)
        placeholders = Placeholder.query.filter_by(template_id=template.id)\
            .order_by(Placeholder.paragraph_index, Placeholder.start_run_index).all()
        
        # Process placeholders (same logic as in generate route)
        for placeholder in placeholders:
            paragraph = doc.paragraphs[placeholder.paragraph_index]
            if placeholder.start_run_index >= len(paragraph.runs) or placeholder.end_run_index >= len(paragraph.runs):
                continue

            user_input = user_inputs.get(placeholder.name, "")
            formatted_text = user_input

            if "date" in placeholder.name.lower() or "date_ofbirth" in placeholder.name.lower():
                formatted_text = format_date(user_input, template.type)

            elif "address" in placeholder.name.lower() and template.type == "letter":
                parts = [part.strip() for part in user_input.split(",")]
                if parts:
                    if placeholder.start_run_index != placeholder.end_run_index:
                        for r_idx in range(placeholder.start_run_index + 1, placeholder.end_run_index + 1):
                            paragraph.runs[r_idx].text = ""
                    run = paragraph.runs[placeholder.start_run_index]
                    run.clear()
                    for i, part in enumerate(parts):
                        run.add_text(part)
                        if i == len(parts) - 1 or part.endswith("."):
                            if not part.endswith("."):
                                run.add_text(".")
                            break
                        else:
                            run.add_text(",")
                            run.add_break()
                    run.font.name = template.font_family
                    run.font.size = Pt(template.font_size)
                    run.bold = placeholder.bold
                    run.italic = placeholder.italic
                    run.underline = placeholder.underline
                    continue

            else:
                if placeholder.casing == "upper":
                    formatted_text = formatted_text.upper()
                elif placeholder.casing == "lower":
                    formatted_text = formatted_text.lower()
                elif placeholder.casing == "title":
                    formatted_text = formatted_text.title()

            run = paragraph.runs[placeholder.start_run_index]
            if placeholder.start_run_index == placeholder.end_run_index:
                run.text = formatted_text
            else:
                for r_idx in range(placeholder.start_run_index + 1, placeholder.end_run_index + 1):
                    paragraph.runs[r_idx].text = ""
                run.text = formatted_text
            run.font.name = template.font_family
            run.font.size = Pt(template.font_size)
            run.bold = placeholder.bold
            run.italic = placeholder.italic
            run.underline = placeholder.underline
        
        # Apply enhanced formatting
        remove_empty_runs(doc)
        enhance_document_formatting(doc, template.type)
        add_page_numbers(doc)
        
        # Save document to disk
        user_name = user_inputs.get("name", "Unknown").strip()
        user_name = re.sub(r'\s+', '_', user_name)
        template_name = template.name.strip()
        template_name = re.sub(r'\s+', '_', template_name)
        current_date = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        file_name = f"{user_name}_{template_name}_{current_date}.docx"
        
        # Save to disk
        file_path = os.path.join(app.config['GENERATED_FOLDER'], file_name)
        doc.save(file_path)
        
        # Create database record with batch_id
        created_doc = CreatedDocument(template_id=template.id, user_name=user_name, file_path=file_name, batch_id=batch_id)
        db.session.add(created_doc)
        generated_docs.append(created_doc)
        successful_generations += 1
    
    db.session.commit()
    
    # Check if any documents were successfully generated
    if successful_generations == 0:
        return render_template('error.html', message="No valid templates were found or all template files are missing. Please contact administrator."), 500
    
    # Redirect to batch results page
    return redirect(url_for('show_batch_results', batch_id=batch_id))

@app.route('/results/<int:doc_id>')
def show_results(doc_id):
    """Display results page for a single generated document."""
    doc = CreatedDocument.query.get_or_404(doc_id)
    return render_template('results.html', document=doc)

@app.route('/batch-results/<batch_id>')
def show_batch_results(batch_id):
    """Display results page for batch generated documents."""
    docs = CreatedDocument.query.filter_by(batch_id=batch_id).all()
    if not docs:
        return render_template('error.html', message="Batch not found or no documents generated."), 404
    return render_template('batch_results.html', documents=docs, batch_id=batch_id)

@app.route('/download-docx/<int:doc_id>')
def download_docx(doc_id):
    """Download a document as DOCX."""
    doc = CreatedDocument.query.get_or_404(doc_id)
    file_path = os.path.join(app.config['GENERATED_FOLDER'], doc.file_path)
    if not os.path.exists(file_path):
        abort(404)
    return send_file(file_path, as_attachment=True)

@app.route('/download-pdf/<int:doc_id>')
def download_pdf(doc_id):
    """Download a document as PDF."""
    doc = CreatedDocument.query.get_or_404(doc_id)
    docx_path = os.path.join(app.config['GENERATED_FOLDER'], doc.file_path)
    if not os.path.exists(docx_path):
        abort(404)
    
    # Generate PDF path
    pdf_filename = doc.file_path.replace('.docx', '.pdf')
    pdf_path = os.path.join(app.config['GENERATED_FOLDER'], pdf_filename)
    
    # Convert to PDF if not exists
    if not os.path.exists(pdf_path):
        success = convert_docx_to_pdf(docx_path, pdf_path)
        if not success:
            return render_template('error.html', message="PDF conversion not available. Please download as DOCX instead."), 500
    
    return send_file(pdf_path, as_attachment=True)

@app.route('/download-all-docx/<batch_id>')
def download_all_docx(batch_id):
    """Download all documents in a batch as DOCX files in ZIP."""
    docs = CreatedDocument.query.filter_by(batch_id=batch_id).all()
    if not docs:
        abort(404)
    
    # Create ZIP file in memory
    memory_file = io.BytesIO()
    with zipfile.ZipFile(memory_file, 'w') as zf:
        for doc in docs:
            file_path = os.path.join(app.config['GENERATED_FOLDER'], doc.file_path)
            if os.path.exists(file_path):
                zf.write(file_path, doc.file_path)
    
    memory_file.seek(0)
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    zip_filename = f"MyTypist_Batch_DOCX_{timestamp}.zip"
    
    return send_file(memory_file, mimetype='application/zip', as_attachment=True, download_name=zip_filename)

@app.route('/download-all-pdf/<batch_id>')
def download_all_pdf(batch_id):
    """Download all documents in a batch as PDF files in ZIP."""
    docs = CreatedDocument.query.filter_by(batch_id=batch_id).all()
    if not docs:
        abort(404)
    
    # Create ZIP file in memory
    memory_file = io.BytesIO()
    with zipfile.ZipFile(memory_file, 'w') as zf:
        for doc in docs:
            docx_path = os.path.join(app.config['GENERATED_FOLDER'], doc.file_path)
            if os.path.exists(docx_path):
                # Generate PDF
                pdf_filename = doc.file_path.replace('.docx', '.pdf')
                pdf_path = os.path.join(app.config['GENERATED_FOLDER'], pdf_filename)
                
                # Convert to PDF if not exists
                if not os.path.exists(pdf_path):
                    convert_docx_to_pdf(docx_path, pdf_path)
                
                # Add PDF to ZIP if conversion was successful
                if os.path.exists(pdf_path):
                    zf.write(pdf_path, pdf_filename)
    
    memory_file.seek(0)
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    zip_filename = f"MyTypist_Batch_PDF_{timestamp}.zip"
    
    return send_file(memory_file, mimetype='application/zip', as_attachment=True, download_name=zip_filename)


@app.route('/download/<int:document_id>')
def download(document_id):
    """Download a previously generated document."""
    doc = CreatedDocument.query.get_or_404(document_id)
    file_path = os.path.join(app.config['GENERATED_FOLDER'], doc.file_path)
    if not os.path.exists(file_path):
        abort(404)
    return send_file(file_path, as_attachment=True)

# **Admin Routes**
@app.route('/admin')
def admin():
    """Display the admin dashboard."""
    key = request.args.get('key')
    if key != app.config['ADMIN_KEY']:
        abort(403)
    templates = Template.query.all()
    total_templates = Template.query.count()
    total_created = CreatedDocument.query.count()
    return render_template('admin.html', templates=templates, total_templates=total_templates,
                         total_created=total_created, admin_key=key)

@app.route('/admin/upload', methods=['POST'])
def upload_template():
    """Upload a new template and extract its placeholders."""
    key = request.form.get('key')
    if key != app.config['ADMIN_KEY']:
        abort(403)
    name = request.form['name']
    type_ = request.form['type']
    file = request.files['file']
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        doc = Document(file_path)
        font_family, font_size = detect_document_font(doc)
        template = Template(name=name, type=type_, file_path=filename,
                          font_family=font_family, font_size=font_size)
        db.session.add(template)
        db.session.commit()
        placeholders = extract_placeholders(doc)
        multi_run_placeholders = [ph for ph in placeholders if ph['start_run_index'] != ph['end_run_index']]
        if multi_run_placeholders:
            names = [ph['name'] for ph in multi_run_placeholders]

        for ph in placeholders:
            placeholder = Placeholder(**ph, template_id=template.id)
            db.session.add(placeholder)
        db.session.commit()
        return redirect(url_for('admin', key=key))
    return "Invalid file", 400

@app.route('/admin/edit/<int:template_id>')
def edit_template(template_id):
    """Render the template edit page."""
    key = request.args.get('key')
    if key != app.config['ADMIN_KEY']:
        abort(403)
    template = Template.query.get_or_404(template_id)
    placeholders = Placeholder.query.filter_by(template_id=template_id).all()
    return render_template('edit.html', template=template, placeholders=placeholders, admin_key=key)

@app.route('/admin/update/<int:template_id>', methods=['POST'])
def update_template(template_id):
    """Update a template's details and placeholder styles."""
    key = request.form.get('key')
    if key != app.config['ADMIN_KEY']:
        abort(403)
    template = Template.query.get_or_404(template_id)
    template.name = request.form['name']
    template.type = request.form['type']
    template.font_family = request.form['font_family']
    template.font_size = int(request.form['font_size'])
    for ph in template.placeholders:
        ph.bold = f'bold_{ph.id}' in request.form
        ph.italic = f'italic_{ph.id}' in request.form
        ph.underline = f'underline_{ph.id}' in request.form
        ph.casing = request.form[f'casing_{ph.id}']
    db.session.commit()
    return redirect(url_for('admin', key=key))

@app.route('/admin/pause/<int:template_id>')
def pause_template(template_id):
    """Pause a template (set is_active to False)."""
    key = request.args.get('key')
    if key != app.config['ADMIN_KEY']:
        abort(403)
    template = Template.query.get_or_404(template_id)
    template.is_active = False
    db.session.commit()
    return redirect(url_for('admin', key=key))

@app.route('/admin/resume/<int:template_id>')
def resume_template(template_id):
    """Resume a paused template (set is_active to True)."""
    key = request.args.get('key')
    if key != app.config['ADMIN_KEY']:
        abort(403)
    template = Template.query.get_or_404(template_id)
    template.is_active = True
    db.session.commit()
    return redirect(url_for('admin', key=key))

@app.route('/delete/<int:document_id>', methods=['GET', 'POST'])
def delete(document_id):
    """Delete a generated document and its file."""
    doc = CreatedDocument.query.get_or_404(document_id)
    file_path = os.path.join(app.config['GENERATED_FOLDER'], doc.file_path)
    if os.path.exists(file_path):
        os.remove(file_path)
    db.session.delete(doc)
    db.session.commit()
    return redirect(url_for('index'))

@app.route('/admin/delete/<int:template_id>')
def delete_template(template_id):
    """Delete a template and its associated data."""
    key = request.args.get('key')
    if key != app.config['ADMIN_KEY']:
        abort(403)
    template = Template.query.get_or_404(template_id)
    db.session.delete(template)
    db.session.commit()
    return redirect(url_for('admin', key=key))

# Run the app locally (not used on PythonAnywhere)
if __name__ == '__main__':
    with app.app_context():
        db.create_all()  # Creates the database tables
    app.run(host='0.0.0.0', port=8000, debug=True)